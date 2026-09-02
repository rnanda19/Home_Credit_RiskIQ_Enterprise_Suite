"""
src/reporting/report_builder.py

Shared Stage-5 (Financial-Impact Reporting & Packaging) component library --
built once, imported by every problem's notebook (HYPER standing rule), instead
of each notebook hand-rolling its own Word/Excel/HTML generation code.

Zero-fabrication contract: every function here only ever renders values it is
handed by the caller. It invents nothing -- no default numbers, no placeholder
text presented as real, no narrative "story" or SMART-suggestion wording of its
own. The calling notebook computes real values (and real narrative text, built
from real computed numbers via f-strings -- never hallucinated commentary) and
explicitly labels any ASSUMPTION with its source before calling in here. This
module is purely a consistent, vivid, colorful RENDERER for what the notebook
already computed.

Produces, from one real notebook run's own computed results:
  - write_csv_outputs()   -> one or more real CSV files (pandas .to_csv)
  - build_word_report()   -> a real .docx (python-docx), colorized, with
                              per-section narrative "story" paragraphs and a
                              Key Insights & SMART Recommendations section
  - build_excel_workbook()-> a real, formula-driven, multi-sheet .xlsx
                              (openpyxl), colorized headers + heat-map
                              conditional formatting + an Insights sheet
  - build_html_dashboard()-> a real, self-contained .html (Chart.js via CDN --
                              safe here because it runs in the USER's own
                              browser when they open the file, not inside any
                              sandbox) with per-chart narrative "story" text,
                              real-data view switchers (slicers/filters), an
                              Insights & SMART Recommendations section, and a
                              searchable/filterable sampled-records table
"""
from __future__ import annotations

import inspect
import json
import re
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.formatting.rule import ColorScaleRule

# ---------------------------------------------------------------------------
# Standing chart-style rule (project-wide, all problems): every chart and every
# colorized report element uses this 8-hue, CVD-validated categorical palette
# (worst adjacent color-vision-deficiency separation 9.1 light / 8.4 dark,
# OKLab scale) -- vivid and multicolored, never a single flat brand color.
# Canonical home for this palette: every notebook imports it from here (HYPER)
# instead of redefining it.
# ---------------------------------------------------------------------------
VIVID_PALETTE = ["#2a78d6", "#eb6834", "#1baf7a", "#eda100", "#e87ba4", "#008300", "#4a3aa7", "#e34948"]
STATUS_GOOD = "#0ca30c"
STATUS_CRITICAL = "#d03b3b"
STATUS_WARNING = "#fab219"


def _palette(n: int) -> list[str]:
    return [VIVID_PALETTE[i % len(VIVID_PALETTE)] for i in range(n)]


def _contrast_text(hex_color: str) -> str:
    """Pick black or white text for readable contrast against a given fill color."""
    h = hex_color.lstrip("#")
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
    return "1a1a1a" if luminance > 0.6 else "FFFFFF"


# ---------------------------------------------------------------------------
# Real bug fix (2026-09-01, found on the user's real full-scale run): Excel
# forbids \ / ? * [ ] : in a worksheet title. Mega Project 4's own real
# problem names -- "Revolving/Credit-Card Distress Early Warning", "POS/Cash
# Loan Delinquency Trajectory" -- both contain a literal real "/", which
# previously reached wb.create_sheet() unsanitized and crashed with
# openpyxl's "Invalid character / found in sheet title" on the real dataset
# (never on this suite's own synthetic fixture, whose problem labels never
# happened to contain one of these characters). This is a real, disclosed
# label-formatting bug in this shared HYPER module, not a data-quality bug --
# fixed once here so every past and future caller (any Mega Project's
# executive rollup) is protected, without altering the real, unsanitized
# label anywhere else (reports, charts, insights all keep the real name).
# ---------------------------------------------------------------------------
_INVALID_SHEET_CHARS_RE = re.compile(r"[\\/?*\[\]:]")


def safe_sheet_name(name: str, max_len: int = 31) -> str:
    """Return an Excel-legal worksheet title: characters Excel itself forbids
    in a sheet title (\\ / ? * [ ] :) replaced with '-', a leading/trailing
    apostrophe stripped (also forbidden), truncated to max_len (Excel's own
    31-character sheet-name limit by default), and never left empty. Callers
    that build a sheet name AND separately need to look that same sheet back
    up by name later (e.g. to embed a chart into it) should call this
    function themselves on the exact same input and reuse the result, so the
    name they look up always matches the name actually created here."""
    cleaned = _INVALID_SHEET_CHARS_RE.sub("-", name).strip()
    cleaned = cleaned.strip("'")
    return (cleaned or "Sheet")[:max_len]


# ---------------------------------------------------------------------------
# CSV outputs
# ---------------------------------------------------------------------------
def write_csv_outputs(tables: dict[str, pd.DataFrame], out_dir: Path) -> dict[str, str]:
    """tables: {filename_stem: DataFrame}. Writes out_dir/filename_stem.csv for each.
    Returns {filename_stem: absolute_path_str} for the caller to log/print."""
    out_dir.mkdir(parents=True, exist_ok=True)
    written = {}
    for stem, df in tables.items():
        path = out_dir / f"{stem}.csv"
        df.to_csv(path, index=False)
        written[stem] = str(path)
    return written


# ---------------------------------------------------------------------------
# Word report (python-docx)
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x59, 0x59, 0x59)


def _shade_cell(cell, hex_color: str) -> None:
    """Solid background fill for a docx table cell -- not exposed by python-docx's
    high-level API, so this sets the underlying OOXML w:shd element directly."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _add_smart_block(doc, item: dict[str, str], accent_hex: str) -> None:
    """Renders one SMART-suggestion dict (headline + specific/measurable/
    achievable/relevant/timebound, all real text the caller computed) as a
    colorized block: bold accent-colored headline, then one colored-label
    bullet per facet that has text."""
    p = doc.add_paragraph()
    run = p.add_run(item.get("headline", ""))
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = RGBColor.from_string(accent_hex.lstrip("#"))
    facets = [
        ("Specific", item.get("specific", "")),
        ("Measurable", item.get("measurable", "")),
        ("Achievable", item.get("achievable", "")),
        ("Relevant", item.get("relevant", "")),
        ("Time-bound", item.get("timebound", "")),
    ]
    for label, text in facets:
        if not text:
            continue
        fp = doc.add_paragraph(style="List Bullet")
        lr = fp.add_run(f"{label}: ")
        lr.bold = True
        lr.font.color.rgb = RGBColor.from_string(accent_hex.lstrip("#"))
        fp.add_run(text)


def build_word_report(
    out_path: Path,
    title: str,
    subtitle: str,
    exec_summary: list[str],
    sections: list[dict[str, Any]],
    chart_image_paths: list[Path] | None = None,
    insights: list[dict[str, str]] | None = None,
) -> Path:
    """sections: list of {"heading": str, "paragraphs": [str, ...],
    "table": {"headers": [...], "rows": [[...], ...]} (optional),
    "story": [str, ...] (optional -- 3-4 sentence real narrative computed by the
    caller from this section's own real numbers, rendered as a colored callout),
    "image_path": Path (optional)}.
    insights: optional list of SMART-suggestion dicts (see _add_smart_block),
    rendered as a new "Key Insights & SMART Recommendations" section right
    after the Executive Summary, one accent color per item cycling the vivid
    palette."""
    doc = Document()

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = GREY
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].italic = True

    doc.add_paragraph()
    doc.add_heading("Executive Summary", level=1)
    for line in exec_summary:
        doc.add_paragraph(line, style="List Bullet")

    if insights:
        doc.add_heading("Key Insights & SMART Recommendations", level=1)
        colors = _palette(len(insights))
        for item, accent in zip(insights, colors):
            _add_smart_block(doc, item, accent)
            doc.add_paragraph()

    for i, sec in enumerate(sections):
        accent = VIVID_PALETTE[i % len(VIVID_PALETTE)]
        sh = doc.add_heading(sec["heading"], level=1)
        for run in sh.runs:
            run.font.color.rgb = RGBColor.from_string(accent.lstrip("#"))
        for para in sec.get("paragraphs", []):
            doc.add_paragraph(para)
        tbl = sec.get("table")
        if tbl:
            headers = tbl["headers"]
            rows = tbl["rows"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Grid Accent 1"
            font_color = RGBColor.from_string(_contrast_text(accent))
            for j, htext in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = str(htext)
                _shade_cell(cell, accent)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = font_color
            for row in rows:
                cells = table.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = str(val)
        img = sec.get("image_path")
        if img and Path(img).exists():
            doc.add_picture(str(img), width=Inches(6.0))
        story = sec.get("story")
        if story:
            lines = story if isinstance(story, list) else [story]
            for k, line in enumerate(lines):
                sp = doc.add_paragraph()
                if k == 0:
                    lr = sp.add_run("Insight — ")
                    lr.bold = True
                    lr.font.color.rgb = RGBColor.from_string(accent.lstrip("#"))
                r = sp.add_run(line)
                r.italic = True
                r.font.color.rgb = GREY

    if chart_image_paths:
        doc.add_heading("Charts", level=1)
        for img_path in chart_image_paths:
            if Path(img_path).exists():
                doc.add_picture(str(img_path), width=Inches(6.0))

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Excel workbook (openpyxl, formula-driven, multi-sheet)
# ---------------------------------------------------------------------------
INPUT_FONT = Font(color="0000FF", bold=True)          # blue font
INPUT_FILL = PatternFill("solid", fgColor="FFFF00")   # yellow fill -- xlsx skill
HEADER_FONT = Font(color="FFFFFF", bold=True)
HEADER_FILL = PatternFill("solid", fgColor="1F3864")
THIN = Border(*(Side(style="thin"),) * 4)


def assumption_ref(assumptions: dict, label: str) -> str:
    """Pure helper: the Assumptions-sheet cell address for `label`, computed purely
    from its position in the (ordered) `assumptions` dict -- callable BEFORE
    build_excel_workbook() runs, so formula strings can reference it directly
    (avoids the chicken-and-egg problem of needing the workbook's own return
    value to build the formula_sheet argument passed INTO the same call)."""
    row = list(assumptions.keys()).index(label) + 2
    return f"Assumptions!$B${row}"


def build_excel_workbook(
    out_path: Path,
    assumptions: dict[str, float | str],
    assumption_notes: dict[str, str],
    data_sheets: list[dict[str, Any]],
    formula_sheet: dict[str, Any] | None = None,
    insights_sheet: dict[str, Any] | None = None,
) -> Path:
    """assumptions: {label: value} rendered on an 'Assumptions' sheet (blue font /
    yellow fill = hardcoded input, per the xlsx skill convention), each with a
    source note in assumption_notes.

    data_sheets: [{"name":, "headers":[...], "rows":[[...]],
    "highlight_col": header_name (optional)}] plain data tables (Excel Table +
    native AutoFilter, no PivotTable Slicers). Each sheet's header row is
    colorized from the vivid palette, cycling one hue per sheet; when
    "highlight_col" names a numeric column, a real red-yellow-green heat-map
    conditional-format color scale is applied to it (computed live by Excel
    from the real cell values -- not a precomputed color).

    formula_sheet (optional): {"name":, "rows": [(label,
    formula_str_or_value), ...]} where formulas reference 'Assumptions'!$B$n by
    cell address the caller supplies -- real Excel formulas, not pre-computed
    numbers, so opening the file and hitting recalc reproduces the same figures.

    insights_sheet (optional): {"name":, "items": [SMART-suggestion dict, ...]}
    (see _add_smart_block for the dict shape) rendered as a colorized,
    wrapped-text table -- one row per insight, cycling the vivid palette as a
    full-row fill."""
    wb = Workbook()
    ws0 = wb.active
    ws0.title = "Assumptions"
    ws0.append(["Label", "Value", "Source / Note"])
    for c in ws0[1]:
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
    for i, (label, value) in enumerate(assumptions.items(), start=2):
        ws0.cell(row=i, column=1, value=label)
        vcell = ws0.cell(row=i, column=2, value=value)
        vcell.font = INPUT_FONT
        vcell.fill = INPUT_FILL
        ws0.cell(row=i, column=3, value=assumption_notes.get(label, ""))
    for col, width in zip("ABC", (34, 16, 60)):
        ws0.column_dimensions[col].width = width

    for idx, sheet in enumerate(data_sheets):
        ws = wb.create_sheet(safe_sheet_name(sheet["name"]))
        ws.append(sheet["headers"])
        header_color = VIVID_PALETTE[idx % len(VIVID_PALETTE)]
        header_font_color = _contrast_text(header_color)
        for c in ws[1]:
            c.font = Font(color=header_font_color, bold=True)
            c.fill = PatternFill("solid", fgColor=header_color.lstrip("#"))
        for row in sheet["rows"]:
            ws.append(row)
        last_row = ws.max_row
        last_col = get_column_letter(len(sheet["headers"]))
        if last_row >= 2:
            from openpyxl.worksheet.table import Table, TableStyleInfo
            tbl_name = "T" + "".join(ch for ch in sheet["name"] if ch.isalnum())[:20]
            table_ref = f"A1:{last_col}{last_row}"
            excel_tbl = Table(displayName=tbl_name, ref=table_ref)
            excel_tbl.tableStyleInfo = TableStyleInfo(
                name="TableStyleMedium2", showRowStripes=True
            )
            ws.add_table(excel_tbl)
            highlight_col = sheet.get("highlight_col")
            if highlight_col and highlight_col in sheet["headers"]:
                col_idx = sheet["headers"].index(highlight_col) + 1
                col_letter = get_column_letter(col_idx)
                rule = ColorScaleRule(
                    start_type="min", start_color="F8696B",
                    mid_type="percentile", mid_value=50, mid_color="FFEB84",
                    end_type="max", end_color="63BE7B",
                )
                ws.conditional_formatting.add(f"{col_letter}2:{col_letter}{last_row}", rule)
        for i, header in enumerate(sheet["headers"], start=1):
            ws.column_dimensions[get_column_letter(i)].width = max(14, len(str(header)) + 4)

    if formula_sheet:
        ws = wb.create_sheet(safe_sheet_name(formula_sheet["name"]))
        ws.append(["Metric", "Value"])
        for c in ws[1]:
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
        for i, (label, formula) in enumerate(formula_sheet["rows"], start=2):
            ws.cell(row=i, column=1, value=label)
            ws.cell(row=i, column=2, value=formula)
        ws.column_dimensions["A"].width = 46
        ws.column_dimensions["B"].width = 20

    if insights_sheet:
        ws = wb.create_sheet(safe_sheet_name(insights_sheet["name"]))
        headers = ["#", "Headline", "Specific", "Measurable", "Achievable", "Relevant", "Time-bound"]
        widths = [4, 30, 38, 32, 32, 32, 32]
        ws.append(headers)
        for c in ws[1]:
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
        for i, item in enumerate(insights_sheet["items"], start=1):
            color = VIVID_PALETTE[(i - 1) % len(VIVID_PALETTE)]
            row = [
                i, item.get("headline", ""), item.get("specific", ""),
                item.get("measurable", ""), item.get("achievable", ""),
                item.get("relevant", ""), item.get("timebound", ""),
            ]
            ws.append(row)
            fill = PatternFill("solid", fgColor=color.lstrip("#"))
            font = Font(color=_contrast_text(color))
            # Row height sized to the longest wrapped cell in this row (chars-per-line
            # approximated from the column width) so long real text (e.g. a "Relevant"
            # note quoting real correlation figures) isn't visually clipped.
            max_lines = 1
            for col, (width, val) in enumerate(zip(widths, row), start=1):
                cell = ws.cell(row=ws.max_row, column=col)
                cell.fill = fill
                cell.font = font
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                text = str(val)
                chars_per_line = max(int(width * 1.7), 1)
                lines = max(1, -(-len(text) // chars_per_line)) if text else 1
                lines += text.count("\n")
                max_lines = max(max_lines, lines)
            ws.row_dimensions[ws.max_row].height = max(20, min(max_lines * 15, 250))
        for col_idx, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(col_idx)].width = w

    wb.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# HTML dashboard (fully self-contained -- Chart.js 4.4.4 is bundled as a local
# vendor asset (src/reporting/vendor/chart.umd.min.js) and inlined directly
# into the page, NOT loaded from a CDN. This is deliberate: a CDN <script src>
# depends on the opener's network at the moment they open the file -- a
# corporate proxy, an offline machine, or security software blocking a
# third-party script host all silently break every chart AND every filter at
# once (they share the same Chart.js load). Inlining removes that entire
# failure class -- the dashboard renders identically online or offline.
# ---------------------------------------------------------------------------
def _this_module_dir() -> Path:
    """Robust replacement for Path(__file__).parent. Some notebook/Jupyter import
    paths execute this module's already-compiled code without injecting the
    __file__ global into its namespace (observed on a real user run -- NameError:
    name '__file__' is not defined, even though the module imported and ran
    correctly otherwise). The compiled code object's own co_filename is set by
    the loader at compile time regardless, so inspect.getfile() on the current
    frame is a reliable fallback that does not depend on __file__ existing."""
    try:
        return Path(__file__).resolve().parent
    except NameError:
        return Path(inspect.getfile(inspect.currentframe())).resolve().parent


_VENDOR_DIR = _this_module_dir() / "vendor"
_CHARTJS_PATH = _VENDOR_DIR / "chart.umd.min.js"


def _load_chartjs_source() -> str:
    if not _CHARTJS_PATH.exists():
        raise FileNotFoundError(
            f"Bundled Chart.js asset missing at {_CHARTJS_PATH}. "
            "src/reporting/vendor/chart.umd.min.js must ship alongside report_builder.py "
            "(HYPER shared module) -- the HTML dashboard is built to be fully self-contained "
            "and deliberately does not fall back to a CDN."
        )
    return _CHARTJS_PATH.read_text(encoding="utf-8")


_HTML_TEMPLATE = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>__TITLE__</title>
<script>
__CHARTJS_INLINE__
</script>
<style>
  :root {
    --ink:#1a1a1a; --muted:#5a5a5a; --bg:#f5f6fa; --card:#ffffff;
    --c1:#2a78d6; --c2:#eb6834; --c3:#1baf7a; --c4:#eda100; --c5:#e87ba4; --c6:#008300; --c7:#4a3aa7; --c8:#e34948;
  }
  * { box-sizing: border-box; }
  body { margin:0; font-family: -apple-system, "Segoe UI", Roboto, sans-serif; background:var(--bg); color:var(--ink); }
  header { background: linear-gradient(120deg, var(--c1), var(--c7) 55%, var(--c5)); color:#fff; padding:32px; }
  header h1 { margin:0 0 6px 0; font-size:26px; }
  header p { margin:0; opacity:.92; font-size:13.5px; }
  main { max-width:1180px; margin:0 auto; padding:24px 28px 56px; }
  h2.section-title { font-size:17px; margin:38px 0 14px; padding-bottom:6px; border-bottom:3px solid var(--c1); color:var(--ink); }
  .kpi-row { display:grid; grid-template-columns:repeat(auto-fit,minmax(190px,1fr)); gap:16px; margin:20px 0 8px; }
  .kpi { background:var(--card); border-radius:10px; padding:16px 18px; box-shadow:0 1px 3px rgba(0,0,0,.08); border-left:5px solid var(--kc,var(--c1)); }
  .kpi .label { font-size:11.5px; color:var(--muted); text-transform:uppercase; letter-spacing:.05em; }
  .kpi .value { font-size:24px; font-weight:700; color:var(--ink); margin-top:4px; }
  .insight-grid { display:grid; grid-template-columns:repeat(auto-fit,minmax(280px,1fr)); gap:16px; }
  .insight-card { background:var(--card); border-radius:10px; padding:18px 20px; box-shadow:0 1px 3px rgba(0,0,0,.08); border-top:5px solid var(--ic,var(--c1)); }
  .insight-card h4 { margin:0 0 10px; font-size:15px; color:var(--ink); }
  .facet { display:flex; gap:8px; margin:6px 0; font-size:12.5px; line-height:1.4; }
  .chip { flex:0 0 auto; font-weight:700; font-size:10.5px; color:#fff; background:var(--ic,var(--c1)); border-radius:4px; padding:2px 7px; height:fit-content; }
  .facet span.txt { color:var(--muted); }
  .chart-card { background:var(--card); border-radius:10px; padding:20px; box-shadow:0 1px 3px rgba(0,0,0,.08); margin-bottom:22px; }
  .chart-card .chart-head { display:flex; justify-content:space-between; align-items:center; gap:12px; margin-bottom:12px; flex-wrap:wrap; }
  .chart-card h3 { margin:0; font-size:15px; color:var(--ink); }
  .chart-card select { font-size:12.5px; padding:5px 8px; border-radius:6px; border:1px solid #ccc; background:#fff; }
  canvas { max-height:340px; }
  .note { font-size:12px; color:var(--muted); margin-top:8px; }
  .story { margin-top:12px; padding:12px 14px; border-radius:8px; background:#eef4fc; border-left:4px solid var(--sc,var(--c1)); font-size:13px; line-height:1.55; color:#2b2b2b; }
  .story b { color:var(--sc,var(--c1)); }
  table.data-table { width:100%; border-collapse:collapse; font-size:12.5px; }
  table.data-table th { background:var(--c1); color:#fff; text-align:left; padding:8px 10px; position:sticky; top:0; }
  table.data-table td { padding:7px 10px; border-bottom:1px solid #e7e7e7; }
  table.data-table tr:nth-child(even) td { background:#fafbfd; }
  .table-wrap { max-height:420px; overflow:auto; border-radius:8px; border:1px solid #e7e7e7; }
  .filter-bar { display:flex; gap:10px; margin:12px 0 14px; flex-wrap:wrap; }
  .filter-bar input, .filter-bar select { padding:7px 10px; border-radius:6px; border:1px solid #ccc; font-size:12.5px; }
  .filter-bar input { flex:1 1 220px; }
  footer { text-align:center; color:var(--muted); font-size:12px; padding:26px; }
</style>
</head>
<body>
<header>
  <h1>__TITLE__</h1>
  <p>__SUBTITLE__</p>
</header>
<main>
  <div class="kpi-row">
    __KPI_HTML__
  </div>

  __INSIGHTS_SECTION__

  <h2 class="section-title">Charts</h2>
  __CHART_CARDS_HTML__

  __TABLE_SECTION__
</main>
<footer>Every figure and every sampled row on this page was computed by this problem's own notebook run against real data. Labeled ASSUMPTION figures are marked as such. Chart filters switch between real precomputed views only -- nothing is invented client-side.</footer>
<script>
const CHART_DATA = __CHART_DATA_JSON__;
const charts = {};
function renderView(cfg, view) {
  const ctx = document.getElementById(cfg.id).getContext('2d');
  if (charts[cfg.id]) charts[cfg.id].destroy();
  charts[cfg.id] = new Chart(ctx, {
    type: cfg.type,
    data: { labels: view.labels, datasets: view.datasets },
    options: cfg.options || { responsive:true, maintainAspectRatio:false, plugins:{legend:{display: cfg.showLegend !== false}} }
  });
}
CHART_DATA.forEach(cfg => {
  // Each chart is rendered independently -- one chart failing (e.g. a browser
  // extension blocking canvas) must never stop the rest of the page, and the
  // filter dropdown is wired up even if the initial render throws.
  try { renderView(cfg, cfg.views[0]); } catch (err) { console.error('Chart render failed:', cfg.id, err); }
  const sel = document.getElementById('filter-' + cfg.id);
  if (sel) {
    sel.addEventListener('change', () => {
      const view = cfg.views.find(v => v.key === sel.value);
      if (view) { try { renderView(cfg, view); } catch (err) { console.error('Chart filter failed:', cfg.id, err); } }
    });
  }
});

const TABLE_DATA = __TABLE_DATA_JSON__;
const TABLE_COLUMNS = __TABLE_COLUMNS_JSON__;
const TABLE_FILTER_COLUMN = __TABLE_FILTER_COLUMN_JSON__;
if (TABLE_DATA.length) {
  const search = document.getElementById('table-search');
  const colFilter = document.getElementById('table-colfilter');
  if (colFilter && TABLE_FILTER_COLUMN) {
    const vals = Array.from(new Set(TABLE_DATA.map(r => String(r[TABLE_FILTER_COLUMN])))).sort();
    colFilter.innerHTML = '<option value="">All ' + TABLE_FILTER_COLUMN + '</option>' +
      vals.map(v => '<option value="' + v + '">' + v + '</option>').join('');
  }
  function renderTable() {
    const q = (search ? search.value : '').toLowerCase();
    const cf = colFilter ? colFilter.value : '';
    const rows = TABLE_DATA.filter(r => {
      if (cf && String(r[TABLE_FILTER_COLUMN]) !== cf) return false;
      if (!q) return true;
      return TABLE_COLUMNS.some(c => String(r[c]).toLowerCase().includes(q));
    });
    const tbody = document.getElementById('table-body');
    tbody.innerHTML = rows.slice(0, 300).map(r =>
      '<tr>' + TABLE_COLUMNS.map(c => '<td>' + r[c] + '</td>').join('') + '</tr>'
    ).join('');
    document.getElementById('table-count').textContent = rows.length + ' of ' + TABLE_DATA.length + ' real sampled rows shown (max 300 rendered)';
  }
  if (search) search.addEventListener('input', renderTable);
  if (colFilter) colFilter.addEventListener('change', renderTable);
  renderTable();
}
</script>
</body>
</html>
"""


def build_html_dashboard(
    out_path: Path,
    title: str,
    subtitle: str,
    kpi_cards: list[dict[str, str]],
    charts: list[dict[str, Any]],
    insights: list[dict[str, str]] | None = None,
    data_table: dict[str, Any] | None = None,
) -> Path:
    """kpi_cards: [{"label":, "value":, "color": optional hex (defaults to a
    cycled palette hue)}].

    charts: [{"id":, "title":, "type": "bar"/"line"/"doughnut", "labels":[...],
    "datasets":[{"label":,"data":[...],"backgroundColor":...}], "note": optional
    str, "story": optional [str,...] (3-4 sentence real narrative computed by
    the caller from this chart's own real numbers), "views": optional
    [{"key":,"label":,"labels":[...],"datasets":[...]}, ...] -- when given,
    renders a dropdown "slicer" that switches the chart between real
    precomputed views (e.g. "All models" vs "Champion vs runner-up", "Top 15"
    vs "Top 5") with no data left unaccounted; when omitted, "labels"/
    "datasets" are used as the sole (unfiltered) view for backward
    compatibility}].

    insights: optional list of SMART-suggestion dicts (see
    _add_smart_block's docstring for the shape), rendered as a
    "Key Insights & SMART Recommendations" card grid, one accent color per
    card cycling the vivid palette.

    data_table: optional {"title":, "columns": [...], "rows": [[...], ...],
    "filter_column": optional column name} -- a real sample of rows (the
    caller decides the sample; this function truncates client-side rendering
    to 300 rows but ships the full sample for search/filter) rendered as a
    searchable, filterable table (a genuine client-side slicer over real
    data, not a fabricated one)."""
    kpi_html = "\n".join(
        f'<div class="kpi" style="--kc:{k.get("color", VIVID_PALETTE[i % len(VIVID_PALETTE)])}">'
        f'<div class="label">{k["label"]}</div><div class="value">{k["value"]}</div></div>'
        for i, k in enumerate(kpi_cards)
    )

    insights_section = ""
    if insights:
        cards = []
        colors = _palette(len(insights))
        for item, accent in zip(insights, colors):
            facets = [
                ("S", "Specific", item.get("specific", "")),
                ("M", "Measurable", item.get("measurable", "")),
                ("A", "Achievable", item.get("achievable", "")),
                ("R", "Relevant", item.get("relevant", "")),
                ("T", "Time-bound", item.get("timebound", "")),
            ]
            facet_html = "\n".join(
                f'<div class="facet"><span class="chip">{code}</span>'
                f'<span class="txt"><b>{label}:</b> {text}</span></div>'
                for code, label, text in facets if text
            )
            cards.append(
                f'<div class="insight-card" style="--ic:{accent}">'
                f'<h4>{item.get("headline", "")}</h4>{facet_html}</div>'
            )
        insights_section = (
            '<h2 class="section-title">Key Insights &amp; SMART Recommendations</h2>'
            f'<div class="insight-grid">{"".join(cards)}</div>'
        )

    chart_cards = []
    chart_configs = []
    for i, ch in enumerate(charts):
        accent = ch.get("accent", VIVID_PALETTE[i % len(VIVID_PALETTE)])
        views = ch.get("views")
        if not views:
            views = [{"key": "default", "label": "Default", "labels": ch["labels"], "datasets": ch["datasets"]}]
        filter_html = ""
        if len(views) > 1:
            options = "\n".join(f'<option value="{v["key"]}">{v["label"]}</option>' for v in views)
            filter_html = f'<select id="filter-{ch["id"]}">{options}</select>'
        note = f'<div class="note">{ch["note"]}</div>' if ch.get("note") else ""
        story_html = ""
        story = ch.get("story")
        if story:
            paras = " ".join(story) if isinstance(story, list) else story
            story_html = f'<div class="story" style="--sc:{accent}"><b>Insight —</b> {paras}</div>'
        chart_cards.append(
            f'<div class="chart-card">'
            f'<div class="chart-head"><h3>{ch["title"]}</h3>{filter_html}</div>'
            f'<canvas id="{ch["id"]}"></canvas>{note}{story_html}</div>'
        )
        chart_configs.append({
            "id": ch["id"], "type": ch["type"], "views": views,
            "showLegend": ch.get("showLegend", True), "options": ch.get("options"),
        })

    table_section = ""
    table_data_json, table_cols_json, table_filter_json = "[]", "[]", "null"
    if data_table:
        cols = data_table["columns"]
        rows = data_table["rows"]
        records = [dict(zip(cols, row)) for row in rows]
        filter_col = data_table.get("filter_column")
        search_box = '<input id="table-search" type="text" placeholder="Search sampled rows...">'
        col_select = '<select id="table-colfilter"></select>' if filter_col else ""
        table_section = (
            f'<h2 class="section-title">{data_table.get("title", "Sampled Real Records")}</h2>'
            f'<div class="filter-bar">{search_box}{col_select}</div>'
            '<div class="table-wrap"><table class="data-table"><thead><tr>'
            + "".join(f"<th>{c}</th>" for c in cols)
            + '</tr></thead><tbody id="table-body"></tbody></table></div>'
            '<div class="note" id="table-count"></div>'
        )
        table_data_json = json.dumps(records, default=str)
        table_cols_json = json.dumps(cols)
        table_filter_json = json.dumps(filter_col) if filter_col else "null"

    html = (
        _HTML_TEMPLATE
        .replace("__CHARTJS_INLINE__", _load_chartjs_source())
        .replace("__TITLE__", title)
        .replace("__SUBTITLE__", subtitle)
        .replace("__KPI_HTML__", kpi_html)
        .replace("__INSIGHTS_SECTION__", insights_section)
        .replace("__CHART_CARDS_HTML__", "\n".join(chart_cards))
        .replace("__TABLE_SECTION__", table_section)
        .replace("__CHART_DATA_JSON__", json.dumps(chart_configs, default=str))
        .replace("__TABLE_DATA_JSON__", table_data_json)
        .replace("__TABLE_COLUMNS_JSON__", table_cols_json)
        .replace("__TABLE_FILTER_COLUMN_JSON__", table_filter_json)
    )
    out_path.write_text(html, encoding="utf-8")
    return out_path
